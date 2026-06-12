"""合成 .docx fixture 构造器（Phase 6 解析器 TDD 用）。

仓库不含真实 SOP 样本（在受控环境外），故用 python-docx + 原始 oxml 构造
确定性、可版本化的 docx 字节流，覆盖解析器需要的全部形态：标准样式标题、
中文同义词样式、零样式编号、内联/独立图片、合并单元格表格、TOC 域、空树。

非 test_*.py，pytest 不收集；仅作 import 辅助。
"""

from __future__ import annotations

import io
import zipfile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree as _et
from PIL import Image

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_V_NS = "urn:schemas-microsoft-com:vml"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def inject_header_part(docx_bytes: bytes, *, header_text: str = "页眉文字") -> bytes:
    """向已构造的 docx 字节流注入一个最小 header1.xml part（含 1 个 <w:p>）。

    不更新 [Content_Types].xml / document.xml.rels —— parser 的 discarded_parts()
    只看 zip 内是否存在该 part + 是否非空，不要求 relationship 完整。
    """
    header_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        b'<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">\n'
        b"    <w:p><w:r><w:t>" + header_text.encode("utf-8") + b"</w:t></w:r></w:p>\n"
        b"</w:hdr>"
    )

    in_buf = io.BytesIO(docx_bytes)
    out_buf = io.BytesIO()
    with (
        zipfile.ZipFile(in_buf, "r") as zin,
        zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout,
    ):
        for item in zin.namelist():
            zout.writestr(item, zin.read(item))
        zout.writestr("word/header1.xml", header_xml)
    return out_buf.getvalue()


def tiny_png(color: tuple[int, int, int] = (200, 30, 30), size: int = 8) -> bytes:
    """生成一张极小 PNG 的字节流。"""
    buf = io.BytesIO()
    Image.new("RGB", (size, size), color).save(buf, format="PNG")
    return buf.getvalue()


class DocxBuilder:
    """链式构造 docx。``build()`` 返回字节流。"""

    def __init__(self) -> None:
        self.doc = Document()

    # ---- 标题 ---------------------------------------------------------- #
    def heading(self, text: str, level: int = 1) -> DocxBuilder:
        """标准 ``Heading N`` 样式标题。"""
        self.doc.add_heading(text, level=level)
        return self

    def styled_heading(
        self, text: str, style_name: str, *, outline_lvl: int | None = None
    ) -> DocxBuilder:
        """自定义命名样式（如「章节标题」）的段落，可选 outlineLvl。"""
        styles = self.doc.styles
        if style_name not in [s.name for s in styles]:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            if outline_lvl is not None:
                ppr = style.element.get_or_add_pPr()
                ol = OxmlElement("w:outlineLvl")
                ol.set(qn("w:val"), str(outline_lvl))
                ppr.append(ol)
        self.doc.add_paragraph(text, style=style_name)
        return self

    # ---- 段落 ---------------------------------------------------------- #
    def para(
        self,
        text: str,
        *,
        bold: bool = False,
        size_pt: float | None = None,
        center: bool = False,
    ) -> DocxBuilder:
        p = self.doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        return self

    def formula_para(self, before: str = "", after: str = "") -> DocxBuilder:
        """段落含一个 OMML 公式（<m:oMathPara><m:oMath>...），可选前后文字。"""
        p = self.doc.add_paragraph()
        if before:
            p.add_run(before)
        omathpara = _et.SubElement(p._p, f"{{{_M_NS}}}oMathPara")
        omath = _et.SubElement(omathpara, f"{{{_M_NS}}}oMath")
        mr = _et.SubElement(omath, f"{{{_M_NS}}}r")
        mt = _et.SubElement(mr, f"{{{_M_NS}}}t")
        mt.text = "x^2"
        if after:
            p.add_run(after)
        return self

    def _graphic_run(self, uri: str, with_image: bool, png: bytes | None = None) -> None:
        """新段落 run 内放 <w:drawing><a:graphic><a:graphicData uri=...>；
        with_image=True 时同 run 再放一张 v:imagedata（模拟 fallback 缓存图）。"""
        _A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
        p = self.doc.add_paragraph()
        run = p.add_run()
        drawing = _et.SubElement(run._r, f"{{{_W_NS}}}drawing")
        graphic = _et.SubElement(drawing, f"{{{_A_NS}}}graphic")
        _et.SubElement(graphic, f"{{{_A_NS}}}graphicData", attrib={"uri": uri})
        if with_image:
            png = png or tiny_png()
            tmp_p = self.doc.add_paragraph()
            tmp_run = tmp_p.add_run()
            tmp_run.add_picture(io.BytesIO(png), width=Pt(20))
            blip = tmp_run._r.find(f".//{{{_A_NS}}}blip")
            rid = blip.get(f"{{{_R_NS}}}embed")
            tmp_p._p.getparent().remove(tmp_p._p)
            pict = _et.SubElement(run._r, f"{{{_W_NS}}}pict")
            shape = _et.SubElement(
                pict, f"{{{_V_NS}}}shape", attrib={"style": "width:20pt;height:20pt"}
            )
            _et.SubElement(shape, f"{{{_V_NS}}}imagedata", attrib={f"{{{_R_NS}}}id": rid})

    def smartart_para(self, with_fallback: bool = False) -> DocxBuilder:
        self._graphic_run("http://schemas.openxmlformats.org/drawingml/2006/diagram", with_fallback)
        return self

    def chart_para(self) -> DocxBuilder:
        self._graphic_run("http://schemas.openxmlformats.org/drawingml/2006/chart", False)
        return self

    def two_charts_one_run(self) -> DocxBuilder:
        """同一 run 内放两个 chart graphicData（均无位图）——验证多图形/run 占位数=2。"""
        _A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
        uri = "http://schemas.openxmlformats.org/drawingml/2006/chart"
        p = self.doc.add_paragraph()
        run = p.add_run()
        for _ in range(2):
            drawing = _et.SubElement(run._r, f"{{{_W_NS}}}drawing")
            graphic = _et.SubElement(drawing, f"{{{_A_NS}}}graphic")
            _et.SubElement(graphic, f"{{{_A_NS}}}graphicData", attrib={"uri": uri})
        return self

    def image_para(self, png: bytes | None = None, *, width_pt: float = 60) -> DocxBuilder:
        """独立成段的图片。"""
        png = png or tiny_png()
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_picture(io.BytesIO(png), width=Pt(width_pt))
        return self

    def heading_with_image(
        self, text: str, level: int = 1, png: bytes | None = None
    ) -> DocxBuilder:
        """标题段落内嵌一张图（模拟锚定在章节标题上的浮动图/logo）。"""
        png = png or tiny_png()
        h = self.doc.add_heading(text, level=level)
        h.add_run().add_picture(io.BytesIO(png), width=Pt(20))
        return self

    def vml_image_para(self, png: bytes | None = None) -> DocxBuilder:
        """段落 run 内嵌一张 VML 老格式图（v:imagedata），同时复用 docx 的图片关系。"""
        png = png or tiny_png()
        # 先用 add_picture 走标准路径注入图片关系（拿到 rid + 落 media part）
        tmp_p = self.doc.add_paragraph()
        tmp_run = tmp_p.add_run()
        tmp_run.add_picture(io.BytesIO(png), width=Pt(20))
        blip = tmp_run._r.find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
        assert blip is not None
        rid = blip.get(f"{{{_R_NS}}}embed")
        # 删除临时段落（关系/媒体已保留）
        tmp_p._p.getparent().remove(tmp_p._p)
        # 构造真正的 VML 段落
        target_p = self.doc.add_paragraph()
        run = target_p.add_run()
        pict = _et.SubElement(run._r, f"{{{_W_NS}}}pict")
        shape = _et.SubElement(
            pict, f"{{{_V_NS}}}shape", attrib={"style": "width:20pt;height:20pt"}
        )
        _et.SubElement(shape, f"{{{_V_NS}}}imagedata", attrib={f"{{{_R_NS}}}id": rid})
        return self

    def textbox_para(self, inner_text: str) -> DocxBuilder:
        """段落 run 内嵌一个 VML 文本框，含一段子段落文字。"""
        p = self.doc.add_paragraph()
        run = p.add_run()
        pict = _et.SubElement(run._r, f"{{{_W_NS}}}pict")
        shape = _et.SubElement(
            pict, f"{{{_V_NS}}}shape", attrib={"style": "width:120pt;height:30pt"}
        )
        tbx = _et.SubElement(shape, f"{{{_V_NS}}}textbox")
        tcontent = _et.SubElement(tbx, f"{{{_W_NS}}}txbxContent")
        inner_p = _et.SubElement(tcontent, f"{{{_W_NS}}}p")
        inner_r = _et.SubElement(inner_p, f"{{{_W_NS}}}r")
        inner_t = _et.SubElement(inner_r, f"{{{_W_NS}}}t")
        inner_t.text = inner_text
        return self

    def textbox_with_image_para(self, inner_text: str, png: bytes | None = None) -> DocxBuilder:
        """文本框内含「文字 + 内联图（a:blip）」的段落——验证 txbx 内图归属内层 block。"""
        png = png or tiny_png()
        # 先借标准 add_picture 注册图片关系
        tmp_p = self.doc.add_paragraph()
        tmp_run = tmp_p.add_run()
        tmp_run.add_picture(io.BytesIO(png), width=Pt(20))
        # 取出 drawing 子树备用（rid 已在 drawing 内的 r:embed 中保留），再删 tmp 段落
        drawing = tmp_run._r.find(f"{{{_W_NS}}}drawing")
        drawing_copy = _et.fromstring(_et.tostring(drawing))
        tmp_p._p.getparent().remove(tmp_p._p)
        # 构造外层段落 → pict → shape → textbox → txbxContent → p(text + drawing)
        outer = self.doc.add_paragraph()
        run = outer.add_run()
        pict = _et.SubElement(run._r, f"{{{_W_NS}}}pict")
        shape = _et.SubElement(
            pict, f"{{{_V_NS}}}shape", attrib={"style": "width:140pt;height:60pt"}
        )
        tbx = _et.SubElement(shape, f"{{{_V_NS}}}textbox")
        tcontent = _et.SubElement(tbx, f"{{{_W_NS}}}txbxContent")
        inner_p = _et.SubElement(tcontent, f"{{{_W_NS}}}p")
        inner_r = _et.SubElement(inner_p, f"{{{_W_NS}}}r")
        inner_t = _et.SubElement(inner_r, f"{{{_W_NS}}}t")
        inner_t.text = inner_text
        inner_r.append(drawing_copy)
        return self

    def textbox_with_sdt_para(self, inner_text: str) -> DocxBuilder:
        """文本框 txbxContent 内含一个 <w:sdt> 包裹的段落 —— 验证 SDT 在 txbx 内也被展开。"""
        p = self.doc.add_paragraph()
        run = p.add_run()
        pict = _et.SubElement(run._r, f"{{{_W_NS}}}pict")
        shape = _et.SubElement(
            pict, f"{{{_V_NS}}}shape", attrib={"style": "width:120pt;height:30pt"}
        )
        tbx = _et.SubElement(shape, f"{{{_V_NS}}}textbox")
        tcontent = _et.SubElement(tbx, f"{{{_W_NS}}}txbxContent")
        sdt = _et.SubElement(tcontent, f"{{{_W_NS}}}sdt")
        sdt_content = _et.SubElement(sdt, f"{{{_W_NS}}}sdtContent")
        inner_p = _et.SubElement(sdt_content, f"{{{_W_NS}}}p")
        inner_r = _et.SubElement(inner_p, f"{{{_W_NS}}}r")
        inner_t = _et.SubElement(inner_r, f"{{{_W_NS}}}t")
        inner_t.text = inner_text
        return self

    def nested_textbox_with_image_para(
        self, outer_text: str, inner_text: str, png: bytes | None = None
    ) -> DocxBuilder:
        """两层嵌套文本框：outer txbx 含一段（outer_text）+ 内层 txbx 含一段（inner_text + 图）。

        验证嵌套 txbx 的图不被外层 txbx 段落双计入 — 图应仅归内层 paragraph block。
        """
        png = png or tiny_png()
        # 借标准 add_picture 注册图片关系，取出 drawing 备用
        tmp_p = self.doc.add_paragraph()
        tmp_run = tmp_p.add_run()
        tmp_run.add_picture(io.BytesIO(png), width=Pt(20))
        drawing = tmp_run._r.find(f"{{{_W_NS}}}drawing")
        drawing_copy = _et.fromstring(_et.tostring(drawing))
        tmp_p._p.getparent().remove(tmp_p._p)
        # 构造：outer_p → pict → shape → textbox → txbxContent_A
        #                                         → p(outer_text)
        #                                         → p[pict → shape → textbox → txbxContent_B
        #                                                                       → p(inner_text + drawing_copy)]
        outer_p = self.doc.add_paragraph()
        outer_run = outer_p.add_run()
        pict_a = _et.SubElement(outer_run._r, f"{{{_W_NS}}}pict")
        shape_a = _et.SubElement(
            pict_a, f"{{{_V_NS}}}shape", attrib={"style": "width:200pt;height:80pt"}
        )
        tbx_a = _et.SubElement(shape_a, f"{{{_V_NS}}}textbox")
        tcontent_a = _et.SubElement(tbx_a, f"{{{_W_NS}}}txbxContent")
        # outer 段（in txbx_A）
        a_p = _et.SubElement(tcontent_a, f"{{{_W_NS}}}p")
        a_r = _et.SubElement(a_p, f"{{{_W_NS}}}r")
        a_t = _et.SubElement(a_r, f"{{{_W_NS}}}t")
        a_t.text = outer_text
        # 包含 inner txbx 的段（仍在 txbx_A 内）
        wrap_p = _et.SubElement(tcontent_a, f"{{{_W_NS}}}p")
        wrap_r = _et.SubElement(wrap_p, f"{{{_W_NS}}}r")
        pict_b = _et.SubElement(wrap_r, f"{{{_W_NS}}}pict")
        shape_b = _et.SubElement(
            pict_b, f"{{{_V_NS}}}shape", attrib={"style": "width:160pt;height:60pt"}
        )
        tbx_b = _et.SubElement(shape_b, f"{{{_V_NS}}}textbox")
        tcontent_b = _et.SubElement(tbx_b, f"{{{_W_NS}}}txbxContent")
        # 内层段（in txbx_B），含图
        b_p = _et.SubElement(tcontent_b, f"{{{_W_NS}}}p")
        b_r = _et.SubElement(b_p, f"{{{_W_NS}}}r")
        b_t = _et.SubElement(b_r, f"{{{_W_NS}}}t")
        b_t.text = inner_text
        b_r.append(drawing_copy)
        return self

    def text_with_image(self, before: str, after: str, png: bytes | None = None) -> DocxBuilder:
        """段中内联图：文字 + 图 + 文字（Q206 整段一个 content 节点）。"""
        png = png or tiny_png()
        p = self.doc.add_paragraph()
        p.add_run(before)
        p.add_run().add_picture(io.BytesIO(png), width=Pt(40))
        p.add_run(after)
        return self

    # ---- 表格 ---------------------------------------------------------- #
    def simple_table(self, data: list[list[str]]) -> DocxBuilder:
        rows = len(data)
        cols = len(data[0]) if rows else 0
        table = self.doc.add_table(rows=rows, cols=cols)
        for r, row in enumerate(data):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
        return self

    def merged_table_with_image(self, png: bytes | None = None) -> DocxBuilder:
        """3x3 表：col0 纵向合并（rowspan=2）、第一行 col1-2 横向合并、单元格内嵌图。"""
        png = png or tiny_png()
        table = self.doc.add_table(rows=3, cols=3)
        # 列合并：(0,1) 与 (0,2)
        table.cell(0, 1).merge(table.cell(0, 2))
        # 行合并：(0,0) 与 (1,0)
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 0).text = "合并左"
        table.cell(0, 1).text = "合并上"
        table.cell(2, 2).text = "尾"
        # 单元格内嵌图
        table.cell(2, 0).paragraphs[0].add_run().add_picture(io.BytesIO(png), width=Pt(30))
        return self

    # ---- TOC 域 -------------------------------------------------------- #
    def toc(self, entries: list[str]) -> DocxBuilder:
        """插入一个 TOC 字段域：begin + instrText(TOC) + 条目 + end。"""
        # begin + instrText
        p = self.doc.add_paragraph()
        r = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r._r.append(fld_begin)
        r2 = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        r2._r.append(instr)
        r3 = p.add_run()
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        r3._r.append(sep)
        # 条目（toc 样式，非 heading）
        for entry in entries:
            self.doc.add_paragraph(entry)
        # end
        pe = self.doc.add_paragraph()
        re = pe.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        re._r.append(fld_end)
        return self

    def section_break(self) -> DocxBuilder:
        """在当前位置插入一个分节符（段落 pPr/sectPr）。"""
        p = self.doc.add_paragraph()
        ppr = p._p.get_or_add_pPr()
        sect = OxmlElement("w:sectPr")
        ppr.append(sect)
        return self

    # ---- 输出 ---------------------------------------------------------- #
    def build(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


# --------------------------------------------------------------------------- #
# 常用预设语料
# --------------------------------------------------------------------------- #
def styled_sop() -> bytes:
    """标准样式 SOP：封面 + TOC + Heading1/2/3 + 段落 + 内联图 + 表格。"""
    return (
        DocxBuilder()
        .para("公司机密文件", center=True, size_pt=22)
        .para("受控副本", center=True)
        .toc(["1 目的\t1", "2 范围\t2"])
        .heading("目的", level=1)
        .para("本程序规定了质量记录的控制要求。")
        .heading("范围", level=1)
        .para("适用于全公司。")
        .heading("职责", level=2)
        .text_with_image("流程见图", "所示。")
        .heading("记录", level=3)
        .simple_table([["编号", "名称"], ["R-01", "记录表"]])
        .build()
    )


def synonym_sop() -> bytes:
    """中文同义词样式 SOP：用「章节标题」样式作 L1（无标准 heading）。"""
    return (
        DocxBuilder()
        .styled_heading("目的", "章节标题")
        .para("正文一。")
        .styled_heading("范围", "章节标题")
        .para("正文二。")
        .build()
    )


def unstyled_numbered_sop() -> bytes:
    """零样式编号 SOP：'1 目的'(粗) / 正文 / '2 范围'(粗) / '2.1 厂内'(粗)。"""
    return (
        DocxBuilder()
        .para("1 目的", bold=True)
        .para("规定记录控制。")
        .para("2 范围", bold=True)
        .para("全公司。")
        .para("2.1 厂内", bold=True)
        .para("厂内适用。")
        .build()
    )


def qms_dunhao_sop() -> bytes:
    """QMS 顿号样式：'1、目的'(粗短=标题) 与 '1、设有消防...'(非粗长=正文)。"""
    return (
        DocxBuilder()
        .para("1、目的", bold=True)
        .para("规定。")
        .para("2、范围", bold=True)
        .para("设有消防设施并按规定维护保养确保完好有效随时可用于灭火。", bold=False)
        .build()
    )


def empty_sop() -> bytes:
    """无任何标题 / 编号的纯段落文档。"""
    return (
        DocxBuilder()
        .para("这是一段普通文字没有任何标题结构。")
        .para("又一段普通文字同样没有结构特征可言。")
        .build()
    )


def manual_toc_sop() -> bytes:
    """手动排版目录 SOP（章节标题样式，无 TOC 域）。

    结构：封面 → 手动目录区（连续章节标题，无正文段）→ 正文区（章节标题 + 正文段）。
    方案C 须跳过目录区的三个「章节标题」，从正文区第一个「目的」（后接正文）开始。
    """
    return (
        DocxBuilder()
        .para("公司机密文件", center=True)
        .section_break()
        .para("目 录")
        .styled_heading("目的", "章节标题")  # 目录区：连续标题，无内容跟随
        .styled_heading("范围", "章节标题")
        .styled_heading("程序", "章节标题")
        .styled_heading("目的", "章节标题")  # 正文区：标题后有正文段
        .para("本程序规定了碘吸附器定期试验要求。")
        .styled_heading("范围", "章节标题")
        .para("适用于一期1–4号机组。")
        .styled_heading("程序", "章节标题")
        .para("操作步骤如下。")
        .build()
    )
